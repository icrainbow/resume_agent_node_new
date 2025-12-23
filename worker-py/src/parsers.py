# parsers.py
import re
import docx
import pypdf
from docx.document import Document as _Document
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

# parsers.py
import sys
import traceback

_PRINTED_ENV = False

def _debug_env_once():
    global _PRINTED_ENV
    if _PRINTED_ENV:
        return
    _PRINTED_ENV = True

    try:
        import docx
        import docx.oxml.parser as p
        print("[env] exe:", sys.executable)
        print("[env] docx:", docx.__file__)
        print("[env] docx_ver:", getattr(docx, "__version__", None))
        print("[env] etree_mod:", p.etree.__name__)
        print("[env] etree_file:", getattr(p.etree, "__file__", None))
    except Exception:
        print("[env] debug failed:\n", traceback.format_exc())

# WordprocessingML namespaces used by python-docx
_W_NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}


def _clean_text(s: str) -> str:
    if not s:
        return ""
    # Normalize newlines and trim trailing spaces on each line
    s = s.replace("\r\n", "\n").replace("\r", "\n")
    s = "\n".join(line.rstrip() for line in s.split("\n"))
    # Collapse 3+ empty lines -> max 2
    s = re.sub(r"\n{3,}", "\n\n", s)
    return s.strip()


def _paragraph_text_with_hyperlinks(p: Paragraph) -> str:
    """
    Extract visible text from a paragraph, INCLUDING hyperlink display text.
    python-docx's p.text can miss parts because hyperlinks live under <w:hyperlink>.
    """
    if p is None:
        return ""

    parts = []

    # Iterate direct children of paragraph XML (<w:p>)
    for child in p._p:
        tag = child.tag

        # Normal run <w:r>
        if tag.endswith("}r"):
            texts = child.xpath(".//w:t", namespaces=_W_NS)
            for t in texts:
                if t.text:
                    parts.append(t.text)

            # Handle explicit line breaks <w:br/> inside runs
            brs = child.xpath(".//w:br", namespaces=_W_NS)
            if brs:
                parts.append("\n")

        # Hyperlink <w:hyperlink> (contains runs)
        elif tag.endswith("}hyperlink"):
            texts = child.xpath(".//w:t", namespaces=_W_NS)
            for t in texts:
                if t.text:
                    parts.append(t.text)

            brs = child.xpath(".//w:br", namespaces=_W_NS)
            if brs:
                parts.append("\n")

        # Some docs use <w:smartTag> etc. which can also contain runs
        else:
            texts = child.xpath(".//w:t", namespaces=_W_NS)
            for t in texts:
                if t.text:
                    parts.append(t.text)

    out = "".join(parts)

    # Preserve bullets visually when Word stores them as numbering but not actual "•"
    # (We cannot reconstruct numbering reliably without deeper parsing,
    # but at least keep paragraph.text if XML extraction ends up empty.)
    if not out.strip():
        out = (p.text or "")

    return out.strip()


def _iter_block_items(parent):
    """
    Yield paragraphs and tables in the order they appear.
    Works for Document body and table cells.
    """
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        # Fallback: try to access ._element if present
        parent_elm = getattr(parent, "_element", None)
        if parent_elm is None:
            return

    for child in parent_elm.iterchildren():
        if child.tag.endswith("}p"):
            yield Paragraph(child, parent)
        elif child.tag.endswith("}tbl"):
            yield Table(child, parent)


def _extract_from_table(tbl: Table) -> str:
    lines = []
    for row in tbl.rows:
        for cell in row.cells:
            # Extract blocks inside each cell in order (paragraphs & nested tables)
            cell_lines = []
            for item in _iter_block_items(cell):
                if isinstance(item, Paragraph):
                    t = _paragraph_text_with_hyperlinks(item)
                    if t:
                        cell_lines.append(t)
                elif isinstance(item, Table):
                    t = _extract_from_table(item)
                    if t:
                        cell_lines.append(t)
            # Join within cell
            cell_text = "\n".join(cell_lines).strip()
            if cell_text:
                lines.append(cell_text)
    return "\n".join(lines).strip()


def _extract_doc_body(doc: docx.Document) -> str:
    lines = []
    for item in _iter_block_items(doc):
        if isinstance(item, Paragraph):
            t = _paragraph_text_with_hyperlinks(item)
            if t:
                lines.append(t)
        elif isinstance(item, Table):
            t = _extract_from_table(item)
            if t:
                lines.append(t)
    return "\n".join(lines).strip()


def _extract_headers_footers(doc: docx.Document) -> str:
    """
    Header/footer often contain name/contact lines in resume templates.
    """
    lines = []

    for section in doc.sections:
        # header
        hdr = section.header
        if hdr:
            for p in hdr.paragraphs:
                t = _paragraph_text_with_hyperlinks(p)
                if t:
                    lines.append(t)
            for tbl in hdr.tables:
                t = _extract_from_table(tbl)
                if t:
                    lines.append(t)

        # footer
        ftr = section.footer
        if ftr:
            for p in ftr.paragraphs:
                t = _paragraph_text_with_hyperlinks(p)
                if t:
                    lines.append(t)
            for tbl in ftr.tables:
                t = _extract_from_table(tbl)
                if t:
                    lines.append(t)

    # De-dup simple repeats (some docs repeat header on every section)
    # Keep order but remove exact duplicates.
    seen = set()
    uniq = []
    for x in lines:
        key = x.strip()
        if key and key not in seen:
            uniq.append(x)
            seen.add(key)

    return "\n".join(uniq).strip()


# parsers.py
import zipfile
import xml.etree.ElementTree as ET

import zipfile
import xml.etree.ElementTree as ET

def parse_docx(file_path: str) -> str:
    """
    More accurate DOCX text extractor.
    - Preserves paragraph boundaries
    - Avoids inserting newlines between runs
    - Handles tabs and line breaks
    """
    try:
        with zipfile.ZipFile(file_path) as z:
            with z.open("word/document.xml") as f:
                tree = ET.parse(f)
                root = tree.getroot()

        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

        paras = []
        for p in root.findall(".//w:p", ns):
            parts = []
            # Iterate through the paragraph’s children in order
            for node in p.iter():
                tag = node.tag

                # Text node
                if tag == f"{{{ns['w']}}}t" and node.text:
                    parts.append(node.text)

                # Tab
                elif tag == f"{{{ns['w']}}}tab":
                    parts.append("\t")

                # Line break / carriage return
                elif tag in (f"{{{ns['w']}}}br", f"{{{ns['w']}}}cr"):
                    parts.append("\n")

            text = "".join(parts).strip()
            if text:
                # Normalize: collapse excessive internal whitespace a bit (optional)
                paras.append(text)

        # Join paragraphs with single newline (or "\n\n" if you prefer)
        return "\n".join(paras).strip()

    except Exception as e:
        print(f"[parse_docx] failed: {e}")
        return ""


def parse_pdf(file_path: str) -> str:
    try:
        with open(file_path, "rb") as f:
            pdf_reader = pypdf.PdfReader(f)
            pages = []
            for page in pdf_reader.pages:
                t = page.extract_text() or ""
                t = t.strip()
                if t:
                    pages.append(t)
            return _clean_text("\n\n".join(pages))
    except Exception as e:
        print(f"Error parsing PDF file: {e}")
        return ""
