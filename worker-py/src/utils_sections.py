# utils_sections.py
# 统一稳定版（去重名 / 无 hardcoding / 严格 schema 驱动）
#
# 核心目标：
# - 每次解析必须使用用户上传的 schema 文件（schema_path）
# - schema 格式：{"version":..., "groups":[...], "sections":[...]}
# - sections 采用 start/end snippet 切片；支持 start/end 为 string 或 list[string]
# - 输出结构：[{id,title,text,parentId,isGroup}, ...]
#
# 额外保留：
# - render_preview_html（worker 侧预览）
# - align_optimized_sections（对齐原文/优化）
# - markdown 合并 + PDF/Word 导出
# - docx 模板就地替换（仅在规则显式给 start/end 时使用）
# - LibreOffice 转 PDF

from __future__ import annotations

from typing import List, Dict, Optional, Tuple, Any

import os, re, json, shutil, subprocess
from dataclasses import dataclass

from docx import Document


# =====================================================
# LibreOffice / soffice
# =====================================================

def _find_soffice() -> str:
    p = shutil.which("soffice")
    if p:
        return p

    mac_path = "/Applications/LibreOffice.app/Contents/MacOS/soffice"
    if os.path.exists(mac_path):
        return mac_path

    raise RuntimeError("未找到 LibreOffice（soffice）。请安装 LibreOffice 或将 soffice 加入 PATH。")


# =====================================================
# 文本规范化（用于匹配/切片稳定性）
# =====================================================

def _normalize_text(s: str) -> str:
    s = (
        (s or "")
        .replace("\r\n", "\n")
        .replace("\r", "\n")
        .replace("\u00A0", " ")  # NBSP
        .replace("–", "-")
        .replace("—", "-")
    )

    # PDF 常见断字：P\nroduct -> Product（仅单字母换行 + 后接小写）
    s = re.sub(r"(?m)(?<=\b[A-Za-z])\n(?=[a-z])", "", s)

    # 数字断裂：202\n2 -> 2022
    s = re.sub(r"(?<=\d)\n(?=\d)", "", s)

    return s

def _collapse_blank_lines(s: str) -> str:
    s = _normalize_text(s)
    s = re.sub(r"\n{3,}", "\n\n", s)
    return s.strip()

def _line_start(text: str, idx: int) -> int:
    if idx <= 0:
        return 0
    p = text.rfind("\n", 0, idx)
    return 0 if p < 0 else p + 1

def _line_end(text: str, idx: int) -> int:
    if idx < 0:
        return 0
    p = text.find("\n", idx)
    return len(text) if p < 0 else p


# =====================================================
# Schema loading & tolerant snippet matching
# =====================================================

def _load_schema_json(obj_or_path: Any) -> Dict[str, Any]:
    """
    Accepts:
      - dict (already parsed schema)
      - str path to .json
    """
    if isinstance(obj_or_path, dict):
        return obj_or_path
    if isinstance(obj_or_path, str):
        with open(obj_or_path, "r", encoding="utf-8") as f:
            return json.load(f)
    raise TypeError("schema must be a dict or a json file path")

def _ensure_list(x: Any) -> List[str]:
    if x is None:
        return []
    if isinstance(x, list):
        return [str(i).strip() for i in x if str(i).strip()]
    if isinstance(x, str):
        s = x.strip()
        return [s] if s else []
    s = str(x).strip()
    return [s] if s else []

def _snippet_to_regex(snippet: str) -> str:
    """
    将 snippet 转成更宽松的 regex：
    - 空格：\s+
    - 短横线：[-–—]\s*
    - 允许末尾轻微标点
    """
    s = (snippet or "").strip()
    if not s:
        return ""

    s = _normalize_text(s)
    s = re.sub(r"\s+", " ", s).strip()

    esc = re.escape(s)
    esc = esc.replace(r"\ ", r"\s+")
    esc = esc.replace(r"\-", r"[-–—]\s*")

    if re.search(r"[A-Za-z0-9]$", s):
        esc = esc + r"[ \t]*[.,;:]?"

    return esc

def _find_first_match_pos(text: str, snippets: List[str], start_at: int = 0) -> Optional[Tuple[int, int]]:
    """
    返回 (start,end) 绝对位置：在 text[start_at:] 里找 snippets 中最早出现的一个。
    """
    t = _normalize_text(text or "")
    start_at = max(0, int(start_at))

    best: Optional[Tuple[int, int]] = None

    for sn in snippets or []:
        sn = (sn or "").strip()
        if not sn:
            continue

        pat = _snippet_to_regex(sn)
        if not pat:
            continue

        try:
            m = re.search(pat, t[start_at:], flags=re.IGNORECASE | re.MULTILINE)
            if not m:
                continue
            s = start_at + m.start()
            e = start_at + m.end()
        except re.error:
            # regex 失败时回退 substring
            idx = t.lower().find(sn.lower(), start_at)
            if idx < 0:
                continue
            s = idx
            e = idx + len(sn)

        if best is None or s < best[0]:
            best = (s, e)

    return best


# =====================================================
# Schema-driven splitting (STRICT: must provide schema_path)
# =====================================================

def _validate_schema_shape(schema: Dict[str, Any]) -> None:
    if not isinstance(schema, dict):
        raise ValueError("schema must be a dict")

    secs = schema.get("sections")
    if not isinstance(secs, list) or not secs:
        raise ValueError("schema must contain non-empty 'sections' list")

    # groups 可选，但你给的 schema 是带 groups 的
    groups = schema.get("groups", [])
    if groups is not None and not isinstance(groups, list):
        raise ValueError("'groups' must be a list if present")

@dataclass
class _Span:
    sec_id: str
    title: str
    parentId: Optional[str]
    start_idx: int
    end_idx: int
    text: str


def _find_heading_idx(raw: str, needle: str, start_from: int = 0) -> int:
    """
    Find needle in raw from start_from, case-insensitive, whitespace-tolerant-ish.
    Keep it conservative: we don't do fuzzy; just normalize spaces.
    """
    if not needle:
        return -1
    hay = raw
    # Case-insensitive search
    idx = hay.lower().find(needle.lower(), start_from)
    return idx


def _mk_loose_header_pattern(header: str) -> re.Pattern:
    """
    Build a regex that matches the header at the beginning of text,
    tolerant to PDF line breaks / spaces between characters.

    Example:
      "Product Owner" can match "P\nroduct Owner" or "P roduct   Owner".
    """
    h = (header or "").strip()
    if not h:
        # never match anything
        return re.compile(r"(?!x)x")

    # Escape each char, then allow arbitrary whitespace/newlines between them.
    # Also allow optional leading whitespace before the header.
    pieces = []
    for ch in h:
        if ch.isspace():
            # collapse any whitespace in header itself
            pieces.append(r"\s+")
        else:
            pieces.append(re.escape(ch) + r"\s*")
    pat = r"^\s*" + "".join(pieces) + r"(?:\s*[:\-–—])?\s*\n+"
    return re.compile(pat, flags=re.IGNORECASE)

def _strip_leading_header(text: str, header: str) -> str:
    """
    Remove a leading header line from the section text.
    If no match, return original text.
    """
    t = (text or "")
    if not t.strip():
        return t

    p = _mk_loose_header_pattern(header)
    t2 = p.sub("", t, count=1)

    # Secondary cleanup: sometimes header is repeated twice or linebreaks are weird.
    # If still starts with the same header in a simpler form, remove once more.
    if t2 != t:
        t2 = t2.lstrip("\n")
    return t2


import re
from typing import Dict, Any, List, Tuple

# -----------------------------
# Helpers (robust for PDF broken headings)
# -----------------------------

def _build_norm_map(raw: str) -> Tuple[str, List[int], List[int]]:
    """
    Normalize by removing ALL whitespace chars, uppercasing.
    Returns:
      norm: normalized string
      norm_to_raw: for each index in norm, the corresponding raw index
      raw_to_norm: for each raw index, the current norm index (monotonic)
    """
    norm_chars: List[str] = []
    norm_to_raw: List[int] = []
    raw_to_norm = [-1] * (len(raw) + 1)

    ni = 0
    for ri, ch in enumerate(raw):
        raw_to_norm[ri] = ni
        if ch.isspace():
            continue
        norm_chars.append(ch.upper())
        norm_to_raw.append(ri)
        ni += 1
    raw_to_norm[len(raw)] = ni

    return "".join(norm_chars), norm_to_raw, raw_to_norm


def _find_heading_idx_fuzzy(raw: str, needle: str, start_from_raw: int, norm_pack=None) -> int:
    """
    Find needle in raw, tolerant to whitespace/newlines inside the heading
    by matching on normalized form (strip all whitespace, uppercase).
    Returns raw index. -1 if not found.
    """
    if not needle:
        return -1

    if norm_pack is None:
        norm, norm_to_raw, raw_to_norm = _build_norm_map(raw)
    else:
        norm, norm_to_raw, raw_to_norm = norm_pack

    n = re.sub(r"\s+", "", needle).upper()
    if not n:
        return -1

    start_from_raw = max(0, min(len(raw), start_from_raw))
    start_norm = raw_to_norm[start_from_raw]

    pos = norm.find(n, max(0, start_norm))
    if pos < 0:
        return -1
    if pos >= len(norm_to_raw):
        return -1
    return norm_to_raw[pos]


def _strip_leading_header_block(text: str, start_marker: str) -> str:
    """
    Remove the section header from the beginning of the section text.

    Works for:
      - Normal headings: "PROFESSIONAL SUMMARY"
      - Broken headings from PDF: "P\\nROFESSIONAL SUMMARY"
      - Headings split across 2-3 lines

    Strategy:
      1) Look at first 3 lines.
      2) If their normalized (whitespace-stripped) prefix contains the start_marker normalized,
         drop those header lines and return the remainder.
    """
    if not text:
        return text

    start_norm = re.sub(r"\s+", "", (start_marker or "")).upper()
    if not start_norm:
        return text.strip()

    lines = text.splitlines()

    # Consider the header possibly spanning 1~3 lines
    for k in (1, 2, 3):
        head = "\n".join(lines[:k])
        head_norm = re.sub(r"\s+", "", head).upper()
        if start_norm in head_norm:
            return "\n".join(lines[k:]).lstrip("\n").strip()

    # Fallback: if first line alone matches (common)
    if lines:
        first_norm = re.sub(r"\s+", "", lines[0]).upper()
        if start_norm in first_norm:
            return "\n".join(lines[1:]).lstrip("\n").strip()

    return text.strip()

import re
from typing import Any

_HEADING_RE = re.compile(r"^[A-Z][A-Z0-9 &/\-]{2,}$")  # e.g. "PROFESSIONAL EXPERIENCE"
_HEADING_RE2 = re.compile(r"^[A-Z][A-Za-z0-9 &/\-]{2,}$")  # looser, e.g. "Core Skills"

def _looks_like_heading(line: str, _unused: object = None) -> bool:
    s = (line or "").strip()
    if not s:
        return False

    # very simple heuristic (够用就行)
    if len(s) <= 2:
        return True

    # all-caps style headings
    letters = [c for c in s if c.isalpha()]
    if letters:
        upper_ratio = sum(1 for c in letters if c.isupper()) / len(letters)
        if upper_ratio >= 0.85 and len(s) <= 60:
            return True

    # common resume headings
    common = (
        "summary", "skills", "experience", "education", "certifications",
        "publications", "patents", "projects"
    )
    low = s.lower()
    if any(k in low for k in common) and len(s) <= 60:
        return True

    return False

# -----------------------------
# Main splitter
# -----------------------------

def split_resume_by_schema(raw: str, schema: Dict[str, Any]) -> List[Dict[str, Any]]:
    if not raw or not raw.strip():
        return []

    # IMPORTANT:
    # - We build group nodes from schema["groups"] (structure only, no text slicing).
    # - We slice text ONLY for leaf items from schema["sections"] where isGroup is False.
    groups = schema.get("groups") or []
    sections = schema.get("sections") or []
    if not isinstance(groups, list) or not isinstance(sections, list):
        return []

    norm_pack = _build_norm_map(raw)

    # -----------------------------
    # 1) Locate group anchors (for scoping searches)
    # -----------------------------
    group_anchor: Dict[str, int] = {}
    for g in groups:
        gid = str(g.get("id", "")).strip()
        title = str(g.get("title", "")).strip()
        if not gid:
            continue
        idx = _find_heading_idx_fuzzy(raw, title, 0, norm_pack) if title else -1
        group_anchor[gid] = idx

    # -----------------------------
    # 2) Extract leaf spans ONLY (do NOT slice groups)
    # -----------------------------
    spans: List[_Span] = []
    for sec in sections:
        sid = str(sec.get("id", "")).strip()
        title = str(sec.get("title", sid)).strip() or sid
        parentId = str(sec.get("parentId")).strip() if sec.get("parentId") else None
        is_group = bool(sec.get("isGroup"))

        if not sid:
            continue

        # ✅ Skip group items in sections list; groups are structural only.
        if is_group:
            continue

        if sid == "1.1":
            print("[dbg] group_anchor[1] =", group_anchor.get("1"))
            ga = group_anchor.get("1", -1)
            if ga >= 0:
                print("[dbg] raw[ga:ga+8] =", repr(raw[ga:ga+8]))

        start_candidates = _get_anchor_candidates(sec, "start")
        end_candidates = _get_anchor_candidates(sec, "end")

        search_from = 0
        if parentId and parentId in group_anchor and group_anchor[parentId] >= 0:
            search_from = group_anchor[parentId]

        # ---- find start ----
        s_idx = -1
        chosen_start = ""

        for cand in start_candidates:
            s_idx = _find_heading_idx_fuzzy(raw, cand, search_from, norm_pack)
            if s_idx >= 0:
                chosen_start = cand
                break

        if s_idx < 0:
            for cand in start_candidates:
                s_idx = _find_heading_idx_fuzzy(raw, cand, 0, norm_pack)
                if s_idx >= 0:
                    chosen_start = cand
                    break

        # ---- fallback: DO NOT DROP LEAF ----
        if s_idx < 0:
            if parentId and group_anchor.get(parentId, -1) >= 0:
                # avoid chopping the first letter of the heading; only skip whitespace
                s_idx = group_anchor[parentId]
                while 0 <= s_idx < len(raw) and raw[s_idx] in "\r\n\t ":
                    s_idx += 1
            else:
                s_idx = 0

        # ---- find end ----
        e_idx = -1
        if end_candidates:
            for cand in end_candidates:
                e_idx = _find_heading_idx_fuzzy(raw, cand, s_idx + 1, norm_pack)
                if e_idx >= 0:
                    break

        if e_idx < 0:
            # try next group boundary
            next_group_pos = min(
                [v for v in group_anchor.values() if v > s_idx],
                default=len(raw),
            )
            e_idx = next_group_pos if next_group_pos > s_idx else len(raw)

        if e_idx <= s_idx:
            e_idx = len(raw)

        text = raw[s_idx:e_idx].strip()

        # ---- FIX: strip header by using ACTUAL first line (prevents "ROFESSIONAL" bug) ----
        if text:
            lines = text.splitlines()
            first_line = (lines[0] or "").strip()
            if first_line and _looks_like_heading(first_line, title):
                j = 1
                while j < len(lines) and not (lines[j] or "").strip():
                    j += 1
                text = "\n".join(lines[j:]).strip()

        spans.append(
            _Span(
                sec_id=sid,
                title=title,
                parentId=parentId,
                start_idx=s_idx,
                end_idx=e_idx,
                text=text,
            )
        )

    # -----------------------------
    # 3) Build group nodes (structure only)
    # -----------------------------
    earliest_child: Dict[str, int] = {}
    for sp in spans:
        if sp.parentId:
            earliest_child[sp.parentId] = min(
                earliest_child.get(sp.parentId, 10**18),
                sp.start_idx,
            )

    group_nodes: List[Tuple[int, Dict[str, Any]]] = []
    for g in groups:
        gid = str(g.get("id", "")).strip()
        title = str(g.get("title", gid)).strip()
        if not gid:
            continue

        pos = min(
            [
                p
                for p in [group_anchor.get(gid, -1), earliest_child.get(gid)]
                if p is not None and p >= 0
            ],
            default=10**18,
        )

        group_nodes.append(
            (pos, {"id": gid, "title": title, "text": "", "parentId": None, "isGroup": True})
        )

    # -----------------------------
    # 4) Merge & sort (group first, leaf after)
    # -----------------------------
    all_nodes: List[Tuple[int, Dict[str, Any]]] = []
    for pos, gnode in group_nodes:
        all_nodes.append((pos, gnode))

    for sp in spans:
        # +1 ensures leaf never collides with group at same anchor
        all_nodes.append(
            (
                sp.start_idx + 1,
                {
                    "id": sp.sec_id,
                    "title": sp.title,
                    "text": sp.text,
                    "parentId": sp.parentId,
                    "isGroup": False,
                },
            )
        )

    all_nodes.sort(key=lambda x: (x[0], 0 if x[1].get("isGroup") else 1))

    # -----------------------------
    # 5) Safe de-dup (id + isGroup)
    # -----------------------------
    seen = set()
    out: List[Dict[str, Any]] = []
    for _, node in all_nodes:
        key = (node.get("id"), bool(node.get("isGroup")))
        if key in seen:
            continue
        seen.add(key)
        out.append(node)

    # -----------------------------
    # 6) Debug prints (dev only, safe)
    # -----------------------------
    if os.environ.get("DEBUG_SPLIT") == "1":
        print(
            "[dbg][split]",
            "total=", len(out),
            "groups=", sum(1 for x in out if x.get("isGroup")),
            "leaves=", sum(1 for x in out if not x.get("isGroup")),
        )
        for i, item in enumerate(out):
            sid_dbg = item.get("id")
            title_dbg = item.get("title")
            text_dbg = item.get("text") or ""
            print("\n==== SPLIT", i, sid_dbg, title_dbg, "len=", len(text_dbg))
            print(text_dbg[:1200])

    return out

def _get_anchor_candidates(sec: Any, which: str) -> List[str]:
    """
    Return anchor candidate strings for a section.

    Expected schema styles:
      1) sec["anchors"] = {"start":[...], "end":[...]}   (your current JSON)
      2) sec["anchor"] / sec["start"] / sec["end"] etc. (tolerant fallback)

    `which` should be "start" or "end".
    """
    if not isinstance(which, str) or which not in ("start", "end"):
        return []

    # sec must be a dict
    if not isinstance(sec, dict):
        return []

    # Primary: sec["anchors"][which]
    anchors = sec.get("anchors")
    if isinstance(anchors, dict):
        v = anchors.get(which)
        return _normalize_candidates(v)

    # Fallbacks (tolerant)
    # e.g. sec["start"] / sec["end"]
    v = sec.get(which)
    if v is not None:
        return _normalize_candidates(v)

    # e.g. sec["anchor"] could store a string/list
    v = sec.get("anchor")
    if v is not None:
        return _normalize_candidates(v)

    return []


def _normalize_candidates(v: Any) -> List[str]:
    """
    Normalize candidate(s) to a list[str], filter empty.
    """
    if v is None:
        return []
    if isinstance(v, str):
        s = v.strip()
        return [s] if s else []
    if isinstance(v, (int, float)):
        s = str(v).strip()
        return [s] if s else []
    if isinstance(v, list):
        out: List[str] = []
        for x in v:
            if x is None:
                continue
            if isinstance(x, str):
                s = x.strip()
                if s:
                    out.append(s)
            elif isinstance(x, (int, float)):
                s = str(x).strip()
                if s:
                    out.append(s)
            else:
                # if someone accidentally put objects inside, stringify lightly
                s = str(x).strip()
                if s and s not in ("{}", "[]"):
                    out.append(s)
        return out
    if isinstance(v, dict):
        # If someone passed {"start":[...],"end":[...]} directly here,
        # pick nothing; the caller should pass sec not anchors dict.
        return []
    return []



def split_resume_by_headlines(raw: str) -> List[Dict[str, Any]]:
    """
    Fallback splitter: mainstream resume headings.
    Returns: List[Dict] with keys: id,title,text,parentId,isGroup

    Strategy:
      1) Normalize lines; keep original text blocks.
      2) Detect headings via common heading set + formatting heuristics.
      3) Slice sections by heading indices.
      4) If no headings found, return single section.
    """
    if not raw or not raw.strip():
        return []

    text = raw.replace("\r\n", "\n").replace("\r", "\n")
    lines = text.split("\n")

    # Common resume headings (extendable)
    # Note: keep uppercase for easier matching after normalization
    HEADINGS = [
        "SUMMARY",
        "PROFESSIONAL SUMMARY",
        "PROFILE",
        "ABOUT",
        "HIGHLIGHTS",
        "CORE COMPETENCIES",
        "SKILLS",
        "TECHNICAL SKILLS",
        "TECHNOLOGIES",
        "EXPERIENCE",
        "WORK EXPERIENCE",
        "PROFESSIONAL EXPERIENCE",
        "EMPLOYMENT HISTORY",
        "CAREER HISTORY",
        "PROJECTS",
        "SELECTED PROJECTS",
        "EDUCATION",
        "CERTIFICATIONS",
        "LICENSES",
        "PUBLICATIONS",
        "PATENTS",
        "AWARDS",
        "ACHIEVEMENTS",
        "VOLUNTEER",
        "VOLUNTEER EXPERIENCE",
        "LEADERSHIP",
        "LANGUAGES",
        "INTERESTS",
        "ADDITIONAL INFORMATION",
    ]

    # Build a normalized set for quick lookup
    heading_set = set([h.upper() for h in HEADINGS])

    def _norm_heading_candidate(s: str) -> str:
        s = (s or "").strip()
        # remove bullet prefix
        s = re.sub(r"^[•\-\*\u2022]+\s*", "", s)
        # remove trailing colon
        s = s[:-1] if s.endswith(":") else s
        # collapse spaces
        s = re.sub(r"\s+", " ", s)
        return s.upper().strip()

    def _looks_like_heading_line(line: str) -> Tuple[bool, str]:
        """
        Heuristic: short-ish line, no too many punctuation, and in heading_set or formatted like heading.
        Returns (is_heading, normalized_title)
        """
        original = (line or "").strip()
        if not original:
            return (False, "")

        # ignore very long lines
        if len(original) > 60:
            return (False, "")

        cand = _norm_heading_candidate(original)
        if not cand:
            return (False, "")

        # direct match on known headings
        if cand in heading_set:
            return (True, cand.title())

        # heuristic match:
        # - mostly uppercase (common in CV headings)
        alpha = re.sub(r"[^A-Za-z]", "", original)
        if alpha:
            upper_ratio = sum(1 for ch in alpha if ch.isupper()) / max(1, len(alpha))
        else:
            upper_ratio = 0.0

        # - ends with colon or looks like a title
        ends_colon = original.endswith(":")
        # - allow "Work Experience" style
        title_case_like = bool(re.match(r"^[A-Z][a-z]+(?:\s+[A-Z][a-z]+){0,4}:?$", original))

        if ends_colon or upper_ratio >= 0.85 or title_case_like:
            # If it matches a partial known heading (e.g. "PROFESSIONAL EXPERIENCE")
            # do a contains match conservatively.
            for h in heading_set:
                if cand == h:
                    return (True, h.title())
            # As last resort, treat the line itself as heading
            # but avoid headings that look like a date range or company line
            if re.search(r"\b(19|20)\d{2}\b", original) and re.search(r"\b-\b", original):
                return (False, "")
            if re.search(r"\b(Singapore|Hong Kong|Shanghai|New York|London)\b", original, flags=re.I):
                # location lines often not headings
                return (False, "")
            return (True, original.rstrip(":").strip())

        return (False, "")

    # Locate headings
    heading_idx: List[Tuple[int, str]] = []
    for i, line in enumerate(lines):
        ok, title = _looks_like_heading_line(line)
        if ok and title:
            heading_idx.append((i, title))

    # De-dup adjacent headings (e.g., blank lines or repeated)
    dedup: List[Tuple[int, str]] = []
    last_i = -999
    last_title = None
    for i, t in heading_idx:
        if i == last_i:
            continue
        if last_title and t.upper() == last_title.upper() and i - last_i <= 2:
            continue
        dedup.append((i, t))
        last_i = i
        last_title = t
    heading_idx = dedup

    def _join_lines(a: int, b: int) -> str:
        chunk = "\n".join(lines[a:b]).strip("\n").strip()
        return chunk

    # If no headings detected, return entire as one section
    if not heading_idx:
        return [{
            "id": "s1",
            "title": "Experience",
            "text": text.strip(),
            "parentId": None,
            "isGroup": False,
        }]

    # Build sections by slicing from each heading to next heading
    sections: List[Dict[str, Any]] = []
    for idx, (start_i, title) in enumerate(heading_idx):
        end_i = heading_idx[idx + 1][0] if idx + 1 < len(heading_idx) else len(lines)

        # content starts after the heading line
        body = _join_lines(start_i + 1, end_i)
        # keep empty section text if needed, but your UI expects text possibly empty
        sid = f"s{idx+1}"
        sections.append({
            "id": sid,
            "title": title,
            "text": body,
            "parentId": None,
            "isGroup": False,
        })

    # If the first heading isn't at top, optionally capture preamble as "Header" (common name+contact)
    first_heading_line = heading_idx[0][0]
    preamble = _join_lines(0, first_heading_line)
    if preamble:
        sections = [{
            "id": "s0",
            "title": "Header",
            "text": preamble,
            "parentId": None,
            "isGroup": False,
        }] + sections

    # Remove truly empty trailing sections (optional, but safe)
    cleaned: List[Dict[str, Any]] = []
    for s in sections:
        title = (s.get("title") or "").strip()
        body = (s.get("text") or "").strip()
        if title and (body or title.lower() == "header"):
            cleaned.append(s)

    return cleaned

def split_resume_by_schema_path(raw: str, schema_path: str) -> List[Dict[str, Any]]:
    schema = _load_schema_json(schema_path)
    return split_resume_by_schema(raw, schema)

# =====================================================
# PUBLIC API: split_resume_into_sections (MUST use schema_path)
# =====================================================

def split_resume_into_sections(text: str, schema_path: Optional[str] = None) -> List[Dict]:
    """
    对外统一入口：
    - 强制 schema_path，不允许 silent fallback
    - 返回结构与前端/worker兼容：[{id,title,text,parentId,isGroup}, ...]
    """
    text = _collapse_blank_lines(text or "")
    if not text:
        return []

    if not schema_path:
        raise ValueError("schema_path is required (strict schema-driven parsing)")

    return split_resume_by_schema_path(text, schema_path)


# =====================================================
# align_optimized_sections（保留：给 worker 侧对齐用）
# =====================================================

def align_optimized_sections(orig_sections: List[Dict], opt_sections: List[Dict]) -> List[Dict]:
    opt_map = {s.get("key"): s for s in (opt_sections or []) if s.get("key")}
    aligned = []
    for orig in (orig_sections or []):
        k = orig.get("key")
        opt = opt_map.get(k, {}) if k else {}
        aligned.append({
            "key": k,
            "title": orig.get("title", ""),
            "original": orig.get("content", ""),
            "optimized": opt.get("content", ""),
        })
    return aligned


# =====================================================
# 合并 + 导出（保留）
# =====================================================

def safe_filename_base(resume_file) -> str:
    if not resume_file:
        return "Resume"
    name = os.path.basename(getattr(resume_file, "name", "") or "Resume")
    base, _ = os.path.splitext(name)
    base = re.sub(r"[^\w\-\.]+", "_", base).strip("_")
    return base or "Resume"

def ensure_outputs_dir() -> str:
    out_dir = os.path.join(os.getcwd(), "outputs")
    os.makedirs(out_dir, exist_ok=True)
    return out_dir

def next_versioned_paths(resume_file):
    """
    <原名>_Optimized_v###.pdf/.docx
    """
    out_dir = ensure_outputs_dir()
    base = safe_filename_base(resume_file)
    prefix = f"{base}_Optimized_v"

    existing = set(os.listdir(out_dir))
    n = 1
    while True:
        ver = f"{n:03d}"
        pdf_name = f"{prefix}{ver}.pdf"
        docx_name = f"{prefix}{ver}.docx"
        if pdf_name not in existing and docx_name not in existing:
            return (
                os.path.join(out_dir, pdf_name),
                os.path.join(out_dir, docx_name),
            )
        n += 1

def build_resume_markdown_from_sections(titles: List[str], left_texts: List[str]) -> str:
    titles = titles or []
    left_texts = left_texts or []
    left_texts = left_texts[:len(titles)]

    parts = []
    for title, body in zip(titles, left_texts):
        body = (body or "").strip()
        if not body:
            continue
        parts.append(f"## {title}\n\n{body}\n")
    return "\n".join(parts).strip()

def create_word_document_from_markdown(md: str, docx_path: str) -> str:
    doc = Document()
    lines = (md or "").splitlines()

    for line in lines:
        s = line.rstrip()
        if not s.strip():
            continue

        if s.startswith("## "):
            doc.add_heading(s[3:].strip(), level=2)
        elif s.startswith("# "):
            doc.add_heading(s[2:].strip(), level=1)
        elif s.lstrip().startswith("- "):
            doc.add_paragraph(s.lstrip()[2:].strip(), style="List Bullet")
        else:
            doc.add_paragraph(s.strip())

    doc.save(docx_path)
    return docx_path

def create_pdf_from_markdown(md: str, pdf_path: str) -> str:
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.units import inch
    from reportlab.pdfgen import canvas

    c = canvas.Canvas(pdf_path, pagesize=LETTER)
    width, height = LETTER

    left = 0.75 * inch
    right = 0.75 * inch
    top = 0.75 * inch
    bottom = 0.75 * inch
    max_width = width - left - right

    y = height - top

    def set_font(size: int, bold: bool = False):
        c.setFont("Helvetica-Bold" if bold else "Helvetica", size)

    def wrap_text(text: str, font_size: int) -> List[str]:
        set_font(font_size, False)
        words = text.split(" ")
        lines_out = []
        cur = ""
        for w in words:
            candidate = (cur + " " + w).strip()
            if not candidate:
                continue
            if c.stringWidth(candidate, "Helvetica", font_size) <= max_width:
                cur = candidate
            else:
                if cur:
                    lines_out.append(cur)
                cur = w
        if cur:
            lines_out.append(cur)
        return lines_out

    def new_page():
        nonlocal y
        c.showPage()
        y = height - top

    lines = (md or "").splitlines()
    for raw in lines:
        s = raw.strip()
        if not s:
            y -= 8
            if y < bottom:
                new_page()
            continue

        if s.startswith("## "):
            y -= 6
            if y < bottom:
                new_page()
            set_font(13, bold=True)
            title = s[3:].strip()
            for tline in wrap_text(title, 13):
                c.drawString(left, y, tline)
                y -= 16
                if y < bottom:
                    new_page()
            y -= 4
            continue

        if s.startswith("# "):
            y -= 8
            if y < bottom:
                new_page()
            set_font(15, bold=True)
            title = s[2:].strip()
            for tline in wrap_text(title, 15):
                c.drawString(left, y, tline)
                y -= 18
                if y < bottom:
                    new_page()
            y -= 6
            continue

        if s.startswith("- "):
            set_font(10, bold=False)
            bullet_text = s[2:].strip()
            wrapped = wrap_text(bullet_text, 10)
            for j, wline in enumerate(wrapped):
                prefix = "• " if j == 0 else "  "
                c.drawString(left, y, prefix + wline)
                y -= 13
                if y < bottom:
                    new_page()
            continue

        set_font(10, bold=False)
        wrapped = wrap_text(s, 10)
        for wline in wrapped:
            c.drawString(left, y, wline)
            y -= 13
            if y < bottom:
                new_page()

    c.save()
    return pdf_path


# =====================================================
# HTML 预览（worker 侧用）
# =====================================================

def render_preview_html(sections):
    html = """
    <style>
    .section { border: 1px solid #e5e7eb; border-radius: 10px; margin-bottom: 20px; padding: 12px; background: #fafafa; }
    .section-title { font-weight: 700; font-size: 16px; margin-bottom: 10px; }
    .cols { display: grid; grid-template-columns: 1fr 1fr; gap: 16px; }
    .col { padding: 10px; border-radius: 8px; white-space: pre-wrap; font-family: ui-monospace, SFMono-Regular, Menlo, Monaco, Consolas; font-size: 13px; }
    .orig { background: #ffffff; border: 1px solid #e5e7eb; }
    .opt { background: #f0f9ff; border: 1px solid #bae6fd; }
    .col-title { font-weight: 600; margin-bottom: 6px; color: #555; }
    </style>
    """
    for sec in sections or []:
        html += f"""
        <div class="section">
            <div class="section-title">{sec.get('title', '')}</div>
            <div class="cols">
                <div class="col orig">
                    <div class="col-title">原始内容</div>
                    {sec.get('raw', '')}
                </div>
                <div class="col opt">
                    <div class="col-title">优化后</div>
                    {sec.get('optimized', '')}
                </div>
            </div>
        </div>
        """
    return html


# =====================================================
# docx 模板就地替换（保留：基于 snippet start/end）
# =====================================================

def _clear_paragraph(p):
    for r in p.runs:
        r.text = ""

def _replace_block_between_markers(doc: Document, start_snippet: str, end_snippet: str, new_text: str):
    paras = doc.paragraphs
    start_i = None
    end_i = None

    for i, p in enumerate(paras):
        if start_i is None and start_snippet and start_snippet in p.text:
            start_i = i
        if start_i is not None and end_snippet and end_snippet in p.text:
            end_i = i
            break

    if start_i is None:
        return False

    if end_snippet and end_i is None:
        end_i = len(paras) - 1

    for j in range(start_i, end_i + 1):
        _clear_paragraph(paras[j])

    lines = (new_text or "").splitlines()
    if not lines:
        return True

    if paras[start_i].runs:
        paras[start_i].runs[0].text = lines[0]
    else:
        paras[start_i].add_run(lines[0])

    cursor = start_i + 1
    for line in lines[1:]:
        if cursor <= end_i:
            paras[cursor].add_run(line)
        else:
            doc.add_paragraph(line)
        cursor += 1

    return True

def create_docx_from_template(template_docx_path: str, sections: List[Dict], left_texts: List[str], out_docx_path: str) -> str:
    """
    说明：
    - 为了“严格 schema 驱动”，模板替换也遵循 sections 的 title->text 映射；
    - 但 docx 内具体定位仍用“start/end snippet”模式（否则太容易误伤排版）。
    - 因此：如果你要模板替换可靠，建议 schema 的 start/end 与 docx 模板文本一致。
    """
    doc = Document(template_docx_path)

    title_to_text: Dict[str, str] = {}
    for i, sec in enumerate(sections):
        title_to_text[sec.get("title", f"Section {i+1}")] = (left_texts[i] if i < len(left_texts) else "")

    # 用 schema 的 start/end 来替换（如果你传进来的 sections 里没有 start/end，就无法定位）
    # 这里保持“安全优先”：只在模板中找到 start_snippet 才替换。
    for title, new_text in title_to_text.items():
        if not (new_text or "").strip():
            continue
        # 你可以在外部调用时提供更精确的 start/end；这里无法自动推断
        # 因此默认不做替换（避免误替换）。
        pass

    doc.save(out_docx_path)
    return out_docx_path


# =====================================================
# LibreOffice 转 PDF（保留）
# =====================================================

def convert_docx_to_pdf_legacy(docx_path: str, pdf_path: str):
    soffice = _find_soffice()
    out_dir = os.path.dirname(pdf_path)

    cmd = [
        soffice,
        "--headless",
        "--nologo",
        "--nofirststartwizard",
        "--norestore",
        "--convert-to", "pdf",
        "--outdir", out_dir,
        docx_path,
    ]
    subprocess.run(cmd, check=True)

    generated_pdf = os.path.join(out_dir, os.path.splitext(os.path.basename(docx_path))[0] + ".pdf")
    if not os.path.exists(generated_pdf):
        raise RuntimeError("LibreOffice 未生成 PDF 文件")

    if os.path.abspath(generated_pdf) != os.path.abspath(pdf_path):
        os.replace(generated_pdf, pdf_path)

def convert_docx_to_pdf(docx_path: str, pdf_path: str):
    if not os.path.exists(docx_path):
        raise FileNotFoundError(f"找不到 docx 文件：{docx_path}")

    soffice = _find_soffice()
    out_dir = os.path.dirname(os.path.abspath(pdf_path))
    os.makedirs(out_dir, exist_ok=True)

    cmd = [
        soffice,
        "--headless",
        "--nologo",
        "--nofirststartwizard",
        "--norestore",
        "--convert-to", "pdf",
        "--outdir", out_dir,
        os.path.abspath(docx_path),
    ]

    try:
        subprocess.run(cmd, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
    except subprocess.CalledProcessError as e:
        raise RuntimeError(
            "LibreOffice 转换失败。\n"
            f"stdout:\n{e.stdout.decode(errors='ignore')}\n"
            f"stderr:\n{e.stderr.decode(errors='ignore')}"
        )

    generated_pdf = os.path.join(out_dir, os.path.splitext(os.path.basename(docx_path))[0] + ".pdf")
    if not os.path.exists(generated_pdf):
        raise RuntimeError("LibreOffice 未生成 PDF 文件，请确认 docx 文件未损坏。")

    if os.path.abspath(generated_pdf) != os.path.abspath(pdf_path):
        if os.path.exists(pdf_path):
            os.remove(pdf_path)
        os.replace(generated_pdf, pdf_path)
